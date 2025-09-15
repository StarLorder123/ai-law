import os
import re
import json
from datetime import datetime
from docx import Document

class LawParser:
    def __init__(self):
        # 更新标题匹配模式，只要求以中华人民共和国开头
        self.title_pattern = re.compile(r'^(?:\（.*?\）|\[.*?\]|\【.*?\】)?中华人民共和国.*?(?:\（.*?\）|\[.*?\]|\【.*?\】)?$')
        # 匹配章节的正则表达式
        self.chapter_pattern = re.compile(r'^第[一二三四五六七八九十百千]+[编章节]')
        # 匹配条款的正则表达式
        self.article_pattern = re.compile(r'^第[一二三四五六七八九十百千]+条')
        # 匹配日期的正则表达式
        self.date_pattern = re.compile(r'(\d{4}年\d{1,2}月\d{1,2}日)')

    def extract_law_info(self, text):
        """提取法律的基本信息"""
        lines = text.split('\n')
        
        # 跳过主席令部分
        start_index = 0
        for i, line in enumerate(lines):
            if '主席令' in line:
                # 继续往下找，直到找到真正的法律标题
                for j in range(i+1, len(lines)):
                    if lines[j].strip() and not any(skip_word in lines[j] for skip_word in ['主席令', '号', '施行', '中华人民共和国主席']):
                        start_index = j
                        break
                break
        
        # 如果没有主席令，就使用第一行
        if start_index == 0:
            title = lines[0].strip()
        else:
            title = lines[start_index].strip()
        
        # 提取前言（括号内的内容）
        preface = ""
        # 从标题后面的文本中查找前言
        remaining_text = '\n'.join(lines[start_index+1:])
        if '（' in remaining_text and '）' in remaining_text:
            start = remaining_text.find('（')
            end = remaining_text.find('）') + 1
            preface = remaining_text[start:end]
        
        return {
            "title": title,
            "preface": preface
        }

    def parse_law_content(self, text):
        """解析法律内容"""
        lines = text.split('\n')
        
        # 跳过主席令部分
        start_index = 0
        for i, line in enumerate(lines):
            if '主席令' in line:
                # 继续往下找，直到找到真正的法律内容
                for j in range(i+1, len(lines)):
                    if lines[j].strip() and not any(skip_word in lines[j] for skip_word in ['主席令', '号', '施行', '中华人民共和国主席']):
                        start_index = j
                        break
                break
        
        # 从实际内容开始处理
        lines = lines[start_index:]
        content = []
        current_section = None  # 用于"编"
        current_chapter = None
        current_articles = []
        
        # 检查是否包含"编"
        has_section = any('编' in line for line in lines if self.chapter_pattern.match(line))
        
        for line in lines[1:]:  # 跳过标题
            line = line.strip()
            if not line:
                continue
                
            # 检查是否是编
            if has_section and self.chapter_pattern.match(line) and '编' in line:
                if current_section:
                    content.append(current_section)
                current_section = {
                    'section-title': line,
                    'section-content': []
                }
                current_chapter = None
                current_articles = []
                
            # 检查是否是章
            elif self.chapter_pattern.match(line):
                if current_chapter:
                    if has_section:
                        current_section['section-content'].append(current_chapter)
                    else:
                        content.append(current_chapter)
                
                current_chapter = {
                    'subtitle': line,
                    'subcontent': []
                }
                current_articles = []
                
            # 检查是否是条款
            elif self.article_pattern.match(line):
                current_articles.append(line)
                if current_chapter:
                    current_chapter['subcontent'] = current_articles
        
        # 处理最后一个章节
        if current_chapter:
            if has_section:
                if current_section:
                    current_section['section-content'].append(current_chapter)
            else:
                content.append(current_chapter)
        
        # 处理最后一个编
        if current_section:
            content.append(current_section)
        
        return content

    def clean_title(self, title):
        """清理标题中的括号内容"""
        # 移除所有类型的括号及其内容
        patterns = [
            r'\（.*?\）',  # 中文括号
            r'\[.*?\]',   # 英文方括号
            r'\【.*?\】'  # 中文方括号
        ]
        cleaned_title = title.strip()
        for pattern in patterns:
            cleaned_title = re.sub(pattern, '', cleaned_title)
        return cleaned_title.strip()

    def is_valid_title(self, title):
        """验证标题是否有效"""
        # 清理括号内容后检查是否以中华人民共和国开头
        cleaned_title = self.clean_title(title)
        return cleaned_title.startswith('中华人民共和国')

    def process_docx(self, file_path):
        """处理单个docx文件"""
        doc = Document(file_path)
        text = '\n'.join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        
        # 获取第一行作为可能的标题
        lines = text.split('\n')
        title = None
        original_title = None
        
        # 在前几行中查找符合格式的标题
        for i in range(min(5, len(lines))):
            current_line = lines[i].strip()
            if self.is_valid_title(current_line):
                original_title = current_line
                title = self.clean_title(current_line)
                break
        
        if not title:
            raise ValueError(f"无法找到有效的法律标题: {file_path}")
        
        # 查找主席令和正文的分界
        start_index = 0
        for i, line in enumerate(lines):
            if "主席令" in line:
                # 找到主席令后的第一个非空行作为正文开始
                for j in range(i + 1, len(lines)):
                    if lines[j].strip() and (title in lines[j] or original_title in lines[j]):
                        start_index = j
                        break
                break
        
        # 构建结果
        result = {
            "ftmc": original_title,  # 保留原始标题（包含括号）
            "preface": self.find_preface(lines[start_index:start_index+5]),
            "content": self.parse_law_content(text)
        }
        
        return result, {"law_name": original_title}

    def find_preface(self, lines):
        """查找并提取序言（通常是括号中的通过日期信息）"""
        preface_pattern = re.compile(r'（.*?通过.*?）')
        for line in lines:
            match = preface_pattern.search(line)
            if match:
                return match.group()
        return ""

def process_law_files(input_folder, output_folder):
    """批量处理法律文件"""
    parser = LawParser()
    
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    # 创建输出文件
    output_file = os.path.join(output_folder, 'laws_output.txt')
    description_file = os.path.join(output_folder, 'laws_description.txt')
    
    with open(output_file, 'w', encoding='utf-8') as f_out, \
         open(description_file, 'w', encoding='utf-8') as f_desc:
        
        for filename in os.listdir(input_folder):
            if filename.endswith('.docx'):
                input_path = os.path.join(input_folder, filename)
                
                try:
                    # 处理文件
                    result, description = parser.process_docx(input_path)
                    
                    # 将法律内容转换为JSON字符串并写入（单行）
                    json_str = json.dumps(result, ensure_ascii=False)
                    f_out.write(json_str + '\n')
                    
                    # 将描述信息写入描述文件
                    f_desc.write(f"{description['law_name']}\n")
                    
                    print(f"成功处理: {filename}")
                    
                except Exception as e:
                    print(f"处理 {filename} 时出错: {str(e)}")

class DocxToMarkdown:
    def __init__(self):
        self.chapter_pattern = re.compile(r'^第[一二三四五六七八九十百千]+[编章节]')
        self.article_pattern = re.compile(r'^第[一二三四五六七八九十百千]+条')

    def convert_to_markdown(self, docx_path, output_folder):
        """将docx文件转换为markdown格式"""
        doc = Document(docx_path)
        
        # 获取文件名（不含扩展名）
        base_name = os.path.splitext(os.path.basename(docx_path))[0]
        output_path = os.path.join(output_folder, f"{base_name}.md")
        
        markdown_content = []
        current_level = 0  # 用于跟踪标题层级
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
                
            # 处理标题
            if self.chapter_pattern.match(text):
                # 判断是编、章还是节
                if '编' in text:
                    markdown_content.append(f"# {text}\n")
                elif '章' in text:
                    markdown_content.append(f"## {text}\n")
                else:  # 节
                    markdown_content.append(f"### {text}\n")
            
            # 处理条款
            elif self.article_pattern.match(text):
                markdown_content.append(f"- {text}\n")
            
            # 处理普通段落
            else:
                # 检查是否是文件标题（通常是第一段）
                if not markdown_content:
                    markdown_content.append(f"# {text}\n\n")
                else:
                    markdown_content.append(f"{text}\n")
        
        # 写入文件
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(''.join(markdown_content))
        
        return output_path

def batch_convert_to_markdown(input_folder, output_folder):
    """批量转换文件夹中的所有docx文件为markdown"""
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    converter = DocxToMarkdown()
    converted_files = []
    
    for filename in os.listdir(input_folder):
        if filename.endswith('.docx'):
            input_path = os.path.join(input_folder, filename)
            try:
                output_path = converter.convert_to_markdown(input_path, output_folder)
                converted_files.append(output_path)
                print(f"成功转换: {filename}")
            except Exception as e:
                print(f"转换 {filename} 时出错: {str(e)}")
    
    return converted_files

# 在主函数中添加转换功能
if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 3:
        print("用法: python splitLaws.py <input_folder> <output_folder> [--to-markdown]")
        sys.exit(1)
    
    input_folder = sys.argv[1]
    output_folder = sys.argv[2]
    
    # 检查是否需要转换为markdown
    if len(sys.argv) > 3 and sys.argv[3] == '--to-markdown':
        print("开始转换文件为Markdown格式...")
        converted_files = batch_convert_to_markdown(input_folder, output_folder)
        print(f"转换完成，共转换 {len(converted_files)} 个文件")
    else:
        # 原有的JSON处理逻辑
        process_law_files(input_folder, output_folder)
