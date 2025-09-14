from docx import Document
import win32com.client
from pdf2docx import Converter
from bs4 import BeautifulSoup
import os

class DocReader:
    @staticmethod
    def read_docx(file_path):
        """
        读取docx文件内容
        Args:
            file_path: docx文件路径
        Returns:
            str: 文档中的所有文本内容
        """
        doc = Document(file_path)
        full_text = []
        
        # 读取所有段落
        for para in doc.paragraphs:
            if para.text.strip():  # 只添加非空段落
                full_text.append(para.text)
        
        # 读取所有表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():  # 只添加非空单元格
                        full_text.append(cell.text)
        
        return '\n'.join(full_text)

    @staticmethod
    def read_doc(file_path):
        """
        读取doc文件内容
        Args:
            file_path: doc文件路径
        Returns:
            str: 文档中的所有文本内容
        """
        # 创建 Word 应用程序对象
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False  # 不显示 Word 窗口

        try:
            # 打开 .doc 文件
            doc = word.Documents.Open(file_path)
            content = doc.Content.Text  # 获取文档内容
            return content
        finally:
            # 确保关闭文件和 Word 应用
            doc.Close()
            word.Quit() 

    @staticmethod
    def pdf_to_docx(pdf_path: str, output_path: str = None) -> str:
        """
        将PDF文件转换为DOCX文件
        Args:
            pdf_path: PDF文件路径
            output_path: 输出的DOCX文件路径，如果不指定则使用相同文件名
        Returns:
            str: 转换后的DOCX文件路径
        """
        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"PDF文件不存在: {pdf_path}")
            
        # 如果未指定输出路径，则使用相同的文件名
        if output_path is None:
            output_path = os.path.splitext(pdf_path)[0] + '.docx'
            
        # 确保输出目录存在
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
        
        try:
            # 创建转换器
            cv = Converter(pdf_path)
            # 转换PDF到DOCX
            cv.convert(output_path)
            # 关闭转换器
            cv.close()
            
            return output_path
        except Exception as e:
            raise Exception(f"PDF转换失败: {str(e)}") 

    @staticmethod
    def read_txt(file_path: str, encoding: str = 'utf-8') -> str:
        """
        读取txt文件内容
        Args:
            file_path: txt文件路径
            encoding: 文件编码，默认为utf-8
        Returns:
            str: 文件中的所有文本内容
        """
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                content = file.read()
            return content
        except UnicodeDecodeError:
            # 如果utf-8解码失败，尝试使用gbk编码
            try:
                with open(file_path, 'r', encoding='gbk') as file:
                    content = file.read()
                return content
            except UnicodeDecodeError as e:
                raise UnicodeDecodeError(f"无法以 {encoding} 或 gbk 编码读取文件，请指定正确的编码格式") from e 

    @staticmethod
    def convert_to_html(file_path: str) -> str:
        """
        将docx文件转换为HTML格式
        
        Args:
            file_path (str): docx文件路径
            
        Returns:
            str: 转换后的HTML字符串
            
        Raises:
            FileNotFoundError: 文件不存在时
            Exception: 其他转换错误
        """
        try:
            doc = Document(file_path)
            soup = BeautifulSoup('<div class="document"></div>', 'html.parser')
            main_div = soup.div
            
            # 处理段落
            for para in doc.paragraphs:
                if para.text.strip():  # 跳过空段落
                    p_tag = soup.new_tag('p')
                    
                    # 处理段落中的样式
                    if para.style.name.startswith('Heading'):
                        # 将 Heading 1-6 转换为对应的 h1-h6
                        level = para.style.name[-1]
                        if level.isdigit() and 1 <= int(level) <= 6:
                            p_tag.name = f'h{level}'
                    
                    # 处理段落中的文本和格式
                    for run in para.runs:
                        span = soup.new_tag('span')
                        span.string = run.text
                        
                        # 添加基本样式
                        if run.bold:
                            span['style'] = 'font-weight: bold;'
                        if run.italic:
                            span['style'] = span.get('style', '') + 'font-style: italic;'
                        if run.underline:
                            span['style'] = span.get('style', '') + 'text-decoration: underline;'
                            
                        p_tag.append(span)
                    
                    main_div.append(p_tag)
            
            # 处理表格
            for table in doc.tables:
                table_tag = soup.new_tag('table')
                table_tag['class'] = 'docx-table'
                table_tag['border'] = '1'
                
                for row in table.rows:
                    tr_tag = soup.new_tag('tr')
                    
                    for cell in row.cells:
                        td_tag = soup.new_tag('td')
                        td_tag.string = cell.text.strip()
                        tr_tag.append(td_tag)
                        
                    table_tag.append(tr_tag)
                    
                main_div.append(table_tag)
            
            # 添加基本的CSS样式
            style_tag = soup.new_tag('style')
            style_tag.string = """
                .document { font-family: Arial, sans-serif; line-height: 1.6; }
                .docx-table { border-collapse: collapse; width: 100%; margin: 10px 0; }
                .docx-table td { padding: 8px; border: 1px solid #ddd; }
            """
            soup.insert(0, style_tag)
            
            return str(soup)
            
        except FileNotFoundError:
            raise FileNotFoundError(f"找不到文件：{file_path}")
        except Exception as e:
            raise Exception(f"转换过程中出现错误：{str(e)}")

    @staticmethod
    def save_as_html(file_path: str, output_path: str = None) -> str:
        """
        将docx文件转换为HTML并保存
        
        Args:
            file_path (str): docx文件路径
            output_path (str, optional): 输出的HTML文件路径，如果不指定则使用相同文件名
            
        Returns:
            str: 保存的HTML文件路径
            
        Raises:
            FileNotFoundError: 文件不存在时
            Exception: 其他转换错误
        """
        try:
            # 如果未指定输出路径，则使用相同的文件名
            if output_path is None:
                output_path = os.path.splitext(file_path)[0] + '.html'
                
            # 确保输出目录存在
            os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
            
            # 转换文件
            html_content = DocReader.convert_to_html(file_path)
            
            # 保存HTML文件
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
                
            return output_path
            
        except Exception as e:
            raise Exception(f"保存HTML文件时出现错误：{str(e)}") 